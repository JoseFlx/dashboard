# dashboard_call_center_clean_fix.py
"""
Dashboard Call Center - Patch: limpeza de nomes + paginação corrigida + geração de relatórios completa.
Substitua seu arquivo atual por este.
"""
import re
import unicodedata
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta, date
import io
import zipfile

# PDF / ReportLab
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.lib.utils import ImageReader

# Pillow
from PIL import Image as PILImage

# Matplotlib
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

# Word
from docx import Document
from docx.shared import Inches, Pt

# ------------------------------
# Config e helpers
# ------------------------------
st.set_page_config(page_title="Dashboard Call Center – Nevoli", layout="wide")
BR_DATE = lambda d: d.strftime('%d/%m/%Y') if isinstance(d, (datetime, date)) else str(d)

def format_timedelta(td):
    if pd.isnull(td):
        return "00:00:00"
    if not isinstance(td, pd.Timedelta):
        try:
            td = pd.to_timedelta(td)
        except Exception:
            return "00:00:00"
    total_seconds = int(td.total_seconds())
    horas = total_seconds // 3600
    minutos = (total_seconds % 3600) // 60
    segundos = total_seconds % 60
    return f"{horas:02d}:{minutos:02d}:{segundos:02d}"

def safe_mean(s):
    if s is None or len(s) == 0:
        return pd.Timedelta(0)
    return s.mean()

def percentis_td(series):
    if series is None or series.empty:
        return {"p50": "00:00:00", "p75": "00:00:00", "p90": "00:00:00", "p95": "00:00:00"}
    try:
        if pd.api.types.is_numeric_dtype(series):
            secs = series.dropna()
        else:
            secs = series.dt.total_seconds().dropna()
    except Exception:
        secs = series.dropna()
    out = {}
    for p in [50, 75, 90, 95]:
        val = np.percentile(secs, p) if len(secs) else 0
        out[f"p{p}"] = format_timedelta(pd.to_timedelta(val, unit="s"))
    return out

def numero_br(x):
    try:
        return f"{int(x):,}".replace(",", ".")
    except Exception:
        return str(x)

# ------------------------------
# limpeza de nomes de agente
# ------------------------------
EMOJI_PATTERN = re.compile(
    "[" 
    "\U0001F600-\U0001F64F"
    "\U0001F300-\U0001F5FF"
    "\U0001F680-\U0001F6FF"
    "\U0001F1E0-\U0001F1FF"
    "\U00002700-\U000027BF"
    "\U0000FE00-\U0000FE0F"
    "]+", flags=re.UNICODE)

def remove_emoji(text: str) -> str:
    if not isinstance(text, str):
        return text
    return EMOJI_PATTERN.sub("", text)

def clean_agent_name(raw):
    if pd.isna(raw):
        return ""
    s = str(raw)
    s = s.replace('\ufeff', '').strip()
    s = unicodedata.normalize('NFKC', s)
    s = remove_emoji(s)
    s = "".join(ch for ch in s if unicodedata.category(ch)[0] != "C")
    s = re.sub(r"[^0-9A-Za-zÀ-ÖØ-öø-ÿ\s\-\.\']", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def is_valid_agent_name(s):
    if not s or pd.isna(s):
        return False
    s = str(s).strip()
    if len(s) <= 1:
        return False
    if not re.search(r"[A-Za-zÀ-ÖØ-öø-ÿ]", s):
        return False
    return True

# ------------------------------
# leitura csv (cache para performance)
# ------------------------------
@st.cache_data(show_spinner=False)
def read_csv_bytes(uploaded_file):
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    return pd.read_csv(uploaded_file, sep=";", encoding="latin-1", low_memory=False)

# ------------------------------
# preparação base dos dados
# ------------------------------
def preparacao_base(df_raw, excluir_bot=True, id_col=None):
    col_map = {
        "Nome do Agente": "Agente",
        "Tempo em Atendimento": "Tempo em Atendimento",
        "Tempo em Espera na Fila": "Tempo em Espera",
        "Tempo em Espera": "Tempo em Espera",
        "Data Inicial": "Data Inicial",
        "Hora Inicial": "Hora Inicial"
    }
    df = df_raw.rename(columns={k:v for k,v in col_map.items() if k in df_raw.columns}).copy()
    df.columns = [c.strip().replace('\ufeff','') for c in df.columns]

    for col in ["Tempo em Atendimento", "Tempo em Espera"]:
        if col in df.columns:
            df[col] = pd.to_timedelta(df[col], errors='coerce')

    if 'Data Inicial' in df.columns:
        df['Data Inicial'] = pd.to_datetime(df['Data Inicial'], errors='coerce')
        df['Data'] = df['Data Inicial'].dt.date
        df['weekday'] = df['Data Inicial'].dt.dayofweek
        df['Dia_Semana'] = df['Data Inicial'].dt.day_name()
        df['Mes_Ano'] = df['Data Inicial'].dt.strftime('%m/%Y')
    else:
        df['Data'] = pd.NaT
        df['Mes_Ano'] = np.nan

    if 'Hora Inicial' in df.columns:
        try:
            df['Hora_Inicio'] = pd.to_datetime(df['Hora Inicial'], errors='coerce').dt.hour
        except Exception:
            if 'Data Inicial' in df.columns:
                df['Hora_Inicio'] = pd.to_datetime(df['Data Inicial'], errors='coerce').dt.hour

    if 'Agente' in df.columns:
        df['Agente'] = df['Agente'].astype(str).apply(lambda x: x.strip())
        df['Agente'] = df['Agente'].apply(clean_agent_name)

    if excluir_bot and 'Agente' in df.columns:
        df = df[df['Agente'].str.upper() != "SYNTESIS - OLIVIA BOT"].copy()

    removed = 0
    used_id_col = None
    if id_col and id_col in df.columns:
        before = len(df)
        df = df.drop_duplicates(subset=[id_col], keep='first')
        removed = before - len(df)
        used_id_col = id_col
    else:
        dup_subset = [c for c in ['Agente','Data Inicial','Hora_Inicio','Tempo em Atendimento','Tempo em Espera'] if c in df.columns]
        if dup_subset:
            before = len(df)
            df = df.drop_duplicates(subset=dup_subset, keep='first')
            removed = before - len(df)
            used_id_col = None

    for c in ['Agente', 'Equipe']:
        if c in df.columns:
            try:
                df[c] = df[c].astype('category')
            except Exception:
                pass

    if 'Tempo em Atendimento' in df.columns:
        df['TMA_s'] = df['Tempo em Atendimento'].dt.total_seconds()
    else:
        df['TMA_s'] = np.nan
    if 'Tempo em Espera' in df.columns:
        df['TME_s'] = df['Tempo em Espera'].dt.total_seconds()
    else:
        df['TME_s'] = np.nan

    if 'Agente' in df.columns:
        valid_mask = df['Agente'].apply(is_valid_agent_name)
        df = df.loc[valid_mask].copy()

    df.attrs['duplicates_removed'] = removed
    df.attrs['dedupe_id_col'] = used_id_col
    return df

# ------------------------------
# split periodos e kpis
# ------------------------------
def split_periodos(df, start_dt, end_dt):
    if 'Data Inicial' not in df.columns:
        return df.copy(), df.copy(), None, None
    mask_now = (df['Data Inicial'].dt.date >= start_dt) & (df['Data Inicial'].dt.date <= end_dt)
    df_now = df.loc[mask_now].copy()
    delta_days = (datetime.combine(end_dt, datetime.max.time()) - datetime.combine(start_dt, datetime.min.time())).days + 1
    prev_end = start_dt - timedelta(days=1)
    prev_start = prev_end - timedelta(days=delta_days - 1)
    mask_prev = (df['Data Inicial'].dt.date >= prev_start) & (df['Data Inicial'].dt.date <= prev_end)
    df_prev = df.loc[mask_prev].copy()
    return df_now, df_prev, prev_start, prev_end

def calc_kpis(df):
    total = len(df)
    if 'TMA_s' in df.columns:
        try:
            tma = pd.to_timedelta(df['TMA_s'].dropna().mean(), unit='s')
        except Exception:
            tma = safe_mean(df.get('Tempo em Atendimento'))
    else:
        tma = safe_mean(df.get('Tempo em Atendimento'))
    if 'TME_s' in df.columns:
        try:
            tme = pd.to_timedelta(df['TME_s'].dropna().mean(), unit='s')
        except Exception:
            tme = safe_mean(df.get('Tempo em Espera'))
    else:
        tme = safe_mean(df.get('Tempo em Espera'))
    return total, tma, tme

# ------------------------------
# tabela ranking
# ------------------------------
@st.cache_data(show_spinner=False)
def tabela_rank(df):
    if 'Agente' not in df.columns or df.empty:
        return pd.DataFrame(columns=["Agente","Qtd_Atendimentos","TMA_td","TME_td","TMA","TME","Participação_%"])
    if 'TMA_s' in df.columns and 'TME_s' in df.columns:
        g = df.groupby('Agente', dropna=False, observed=True).agg(
            Qtd_Atendimentos=('Agente','size'),
            TMA_s_mean=('TMA_s','mean'),
            TME_s_mean=('TME_s','mean')
        ).reset_index()
        g['TMA_td'] = pd.to_timedelta(g['TMA_s_mean'].fillna(0), unit='s')
        g['TME_td'] = pd.to_timedelta(g['TME_s_mean'].fillna(0), unit='s')
    else:
        g = df.groupby('Agente', dropna=False, observed=True).agg(
            Qtd_Atendimentos=('Agente','size'),
            TMA_td=('Tempo em Atendimento', lambda s: s.mean() if (s.notna().any()) else pd.Timedelta(0)),
            TME_td=('Tempo em Espera', lambda s: s.mean() if (s.notna().any()) else pd.Timedelta(0))
        ).reset_index()
    g['TMA'] = g['TMA_td'].apply(format_timedelta)
    g['TME'] = g['TME_td'].apply(format_timedelta)
    total = g['Qtd_Atendimentos'].sum()
    g['Participação_%'] = (g['Qtd_Atendimentos'] / total * 100).round(1) if total > 0 else 0.0
    return g

def detectar_outliers_iqr(df_rank, col_td, k=1.5):
    if df_rank.empty or df_rank[col_td].isna().all():
        return set()
    s = df_rank[col_td].dt.total_seconds()
    q1 = s.quantile(0.25)
    q3 = s.quantile(0.75)
    iqr = q3 - q1
    limite_sup = q3 + k * iqr
    mask = s > limite_sup
    return set(df_rank.loc[mask, "Agente"].tolist())

# ------------------------------
# curva de servico e graficos
# ------------------------------
def curva_servico(df, thresholds_s):
    if "Tempo em Espera" not in df.columns or df.empty:
        return [[t, 0.0] for t in thresholds_s]
    if 'TME_s' in df.columns:
        secs = df['TME_s']
    else:
        secs = df["Tempo em Espera"].dt.total_seconds()
    total = (secs.notna()).sum()
    out = []
    for t in thresholds_s:
        pct = ((secs <= t).mean() * 100) if total > 0 else 0.0
        out.append([t, round(pct, 1)])
    return out

def plot_curva_servico(df_now, sla_thresholds):
    curva = curva_servico(df_now, sla_thresholds)
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=[c[0] for c in curva],
                             y=[c[1] for c in curva],
                             mode='lines+markers',
                             marker=dict(size=8),
                             line=dict(width=2)))
    fig.update_layout(title="Curva de Serviço (SLA de Espera)",
                      xaxis_title="Alvo (segundos)",
                      yaxis_title="% ≤ alvo",
                      yaxis=dict(range=[0, 100]),
                      template="plotly_white",
                      height=340)
    return fig

def plot_heatmap(df):
    if "Data Inicial" not in df.columns or "Hora_Inicio" not in df.columns or df.empty:
        return None
    tmp = df.copy()
    tmp["weekday"] = tmp["Data Inicial"].dt.dayofweek
    mat = tmp.pivot_table(index="weekday", columns="Hora_Inicio", values="Agente", aggfunc="count", fill_value=0)
    for h in range(24):
        if h not in mat.columns:
            mat[h] = 0
    mat = mat.reindex(columns=sorted(mat.columns))
    fig = px.imshow(mat, text_auto=False, color_continuous_scale='Blues', aspect="auto",
                    labels=dict(x="Hora", y="Dia da Semana", color="Atendimentos"))
    fig.update_yaxes(tickvals=list(range(7)), ticktext=["Seg","Ter","Qua","Qui","Sex","Sáb","Dom"])
    fig.update_layout(title="Heatmap – Volume (Dia × Hora)", height=380)
    return fig

# ------------------------------
# imagens para PDF/Word (mantive as funções completas)
# ------------------------------
def fig_curva_servico_image(curva):
    fig, ax = plt.subplots(figsize=(6.5, 3.2), dpi=150)
    xs = [c[0] for c in curva]
    ys = [c[1] for c in curva]
    ax.plot(xs, ys, marker='o', linewidth=2)
    ax.set_title("Curva de Serviço (SLA de Espera)")
    ax.set_xlabel("Alvo (segundos)")
    ax.set_ylabel("% ≤ alvo")
    ax.set_ylim(0, 100)
    ax.yaxis.set_major_formatter(mticker.PercentFormatter())
    ax.grid(True, alpha=0.3)
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png', dpi=140)
    plt.close(fig)
    buf.seek(0)
    return buf

def fig_heatmap_volume_image(df):
    if "Data Inicial" not in df.columns or "Hora_Inicio" not in df.columns or df.empty:
        return None
    tmp = df.copy()
    tmp["weekday"] = tmp["Data Inicial"].dt.dayofweek
    mat = tmp.pivot_table(index="weekday", columns="Hora_Inicio", values="Agente", aggfunc="count", fill_value=0)
    for h in range(24):
        if h not in mat.columns:
            mat[h] = 0
    mat = mat.reindex(columns=sorted(mat.columns))
    y_labels = [ {0:'Seg',1:'Ter',2:'Qua',3:'Qui',4:'Sex',5:'Sáb',6:'Dom'}.get(i,str(i)) for i in mat.index ]
    ncols = mat.shape[1]
    nrows = mat.shape[0]
    width_inches = max(6.0, 0.25 * ncols + 3.0)
    height_inches = max(3.2, 0.6 * nrows + 1.0)
    fig, ax = plt.subplots(figsize=(width_inches, height_inches), dpi=140)
    im = ax.imshow(mat.values, aspect='auto', cmap='Blues', origin='lower')
    if ncols <= 12:
        step = 1
    else:
        step = max(1, ncols // 12)
    xticks = list(range(0, ncols, step))
    ax.set_xticks(xticks)
    col_hours = list(mat.columns)
    xtick_labels = [str(col_hours[i]) for i in xticks]
    ax.set_xticklabels(xtick_labels, rotation=45, ha='right', fontsize=8)
    ax.set_yticks(range(len(mat.index)))
    ax.set_yticklabels(y_labels, fontsize=9)
    ax.set_title("Heatmap – Volume (Dia × Hora)")
    ax.set_xlabel("Hora")
    ax.set_ylabel("Dia")
    cbar = plt.colorbar(im, ax=ax, fraction=0.04, pad=0.02)
    cbar.ax.tick_params(labelsize=8)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=140, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf

# ------------------------------
# Custom PDF canvas (rodapé)
# ------------------------------
class NumberedCanvas(pdfcanvas.Canvas):
    logo_bytes = None
    signature_text = None
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []
        self._logo = None
        if NumberedCanvas.logo_bytes:
            try:
                buf = io.BytesIO(NumberedCanvas.logo_bytes)
                img = PILImage.open(buf); img.verify(); buf.seek(0)
                self._logo = ImageReader(buf)
            except Exception:
                self._logo = None

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_footer_and_header(num_pages)
            super().showPage()
        super().save()

    def _draw_footer_and_header(self, page_count):
        width, height = self._pagesize
        if self._logo:
            try:
                iw = 36*mm; ih = 10*mm
                self.drawImage(self._logo, 14*mm, height - 18*mm, width=iw, height=ih, preserveAspectRatio=True, mask='auto')
            except Exception:
                pass
        self.setFont("Helvetica", 8)
        gen = datetime.now().strftime("%d/%m/%Y %H:%M")
        self.drawString(14*mm, 10*mm, f"Emitido em: {gen}")
        page_no = self._pageNumber
        self.drawRightString(width - 14*mm, 10*mm, f"Página {page_no} de {page_count}")
        if page_no == page_count and NumberedCanvas.signature_text:
            try:
                self.setFont("Helvetica-Oblique", 9)
                self.drawCentredString(width / 2.0, 18*mm, f"Assinado por: {NumberedCanvas.signature_text}")
            except Exception:
                pass

# ------------------------------
# Geradores PDF / Word / Excel / ZIP (cópia completa)
# ------------------------------
def gerar_pdf_relatorio(df, inicio_dt, fim_dt, thresholds_s, logo_bytes=None, signature_text=None):
    NumberedCanvas.logo_bytes = logo_bytes
    NumberedCanvas.signature_text = signature_text
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=18*mm, leftMargin=14*mm, rightMargin=14*mm)
    styles = getSampleStyleSheet(); styles.add(ParagraphStyle(name='Small', fontSize=9, leading=11))
    elements = []

    elements.append(Paragraph("RELATÓRIO – ATENDIMENTO (CALL CENTER)", styles['Title']))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph("Empresa: Nevoli Telecom | Suporte Técnico / Atendimento", styles['Normal']))
    elements.append(Paragraph(f"Período analisado: {BR_DATE(inicio_dt)} a {BR_DATE(fim_dt)}", styles['Normal']))
    elements.append(Spacer(1, 8))

    total, tma, tme = calc_kpis(df)
    tma_p = percentis_td(df["Tempo em Atendimento"]) if "Tempo em Atendimento" in df.columns else {}
    tme_p = percentis_td(df["Tempo em Espera"]) if "Tempo em Espera" in df.columns else {}

    kpi_data = [
        ["Indicador", "Valor"],
        ["Total de Atendimentos", numero_br(total)],
        ["TMA Médio", format_timedelta(tma)],
        ["TME Médio (Espera)", format_timedelta(tme)],
        ["TMA p50/p75/p90/p95", f"{tma_p.get('p50','-')} / {tma_p.get('p75','-')} / {tma_p.get('p90','-')} / {tma_p.get('p95','-')}"],
        ["TME p50/p75/p90/p95", f"{tme_p.get('p50','-')} / {tme_p.get('p75','-')} / {tme_p.get('p90','-')} / {tme_p.get('p95','-')}"],
    ]
    kpi_table = Table(kpi_data, hAlign="LEFT", colWidths=[80*mm, 90*mm])
    kpi_table.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),0.5,colors.black),('FONTNAME',(0,0),(-1,0),'Helvetica-Bold')]))
    elements.append(kpi_table)
    elements.append(Spacer(1, 8))

    # Curva
    curva = curva_servico(df, thresholds_s)
    elements.append(Paragraph("Curva de Serviço (SLA de Espera)", styles['Heading2']))
    table_curva = Table([["Alvo (s)", "% ≤ alvo"]] + [[c[0], f"{c[1]:.1f}%"] for c in curva], hAlign="LEFT", colWidths=[40*mm, 40*mm])
    table_curva.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),0.5,colors.black),('FONTNAME',(0,0),(-1,0),'Helvetica-Bold')]))
    elements.append(table_curva)
    elements.append(Spacer(1, 6))
    img_buf = fig_curva_servico_image(curva)
    elements.append(RLImage(img_buf, width=160*mm, height=80*mm))
    elements.append(Spacer(1, 8))

    # KPIs por agente segmentadas: Atendimentos, TMA, TME, Participacao
    rk = tabela_rank(df).sort_values("Qtd_Atendimentos", ascending=False).reset_index(drop=True)
    elements.append(Paragraph("KPIs por Agente (segmentado)", styles['Heading2']))
    elements.append(Spacer(1, 4))

    # Atendimentos
    elements.append(Paragraph("Atendimentos (maior → menor)", styles['Heading3']))
    if not rk.empty:
        data_att = [["Agente","Qtd Atendimentos"]] + rk[["Agente","Qtd_Atendimentos"]].values.tolist()
        t_att = Table(data_att, hAlign="LEFT", colWidths=[100*mm, 40*mm])
        t_att.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.4,colors.black),('BACKGROUND',(0,0),(-1,0),colors.lightgrey)]))
        elements.append(t_att)
    else:
        elements.append(Paragraph("Sem dados.", styles['Small']))
    elements.append(Spacer(1,6))

    # TMA
    elements.append(Paragraph("TMA (médio) — ordenado (menor → maior)", styles['Heading3']))
    if not rk.empty:
        rk_tma = rk.sort_values("TMA_td", ascending=True)
        data_tma = [["Agente","TMA"]] + rk_tma[["Agente","TMA"]].values.tolist()
        t_tma = Table(data_tma, hAlign="LEFT", colWidths=[100*mm,40*mm])
        t_tma.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.4,colors.black),('BACKGROUND',(0,0),(-1,0),colors.lightgrey)]))
        elements.append(t_tma)
    else:
        elements.append(Paragraph("Sem dados.", styles['Small']))
    elements.append(Spacer(1,6))

    # TME
    elements.append(Paragraph("TME (médio) — ordenado (menor → maior)", styles['Heading3']))
    if not rk.empty:
        rk_tme = rk.sort_values("TME_td", ascending=True)
        data_tme = [["Agente","TME"]] + rk_tme[["Agente","TME"]].values.tolist()
        t_tme = Table(data_tme, hAlign="LEFT", colWidths=[100*mm,40*mm])
        t_tme.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.4,colors.black),('BACKGROUND',(0,0),(-1,0),colors.lightgrey)]))
        elements.append(t_tme)
    else:
        elements.append(Paragraph("Sem dados.", styles['Small']))
    elements.append(Spacer(1,6))

    # Participação
    elements.append(Paragraph("Participação (%)", styles['Heading3']))
    if not rk.empty:
        data_part = [["Agente","Participação (%)"]] + [[r[0], f"{r[1]:.1f}%"] for r in rk[["Agente","Participação_%"]].values.tolist()]
        t_part = Table(data_part, hAlign="LEFT", colWidths=[100*mm,40*mm])
        t_part.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.4,colors.black),('BACKGROUND',(0,0),(-1,0),colors.lightgrey)]))
        elements.append(t_part)
    else:
        elements.append(Paragraph("Sem dados.", styles['Small']))
    elements.append(PageBreak())

    # Ranking completo
    elements.append(Paragraph("Ranking por Agente (completo)", styles['Heading2']))
    if not rk.empty:
        data_rank = [["Agente","Qtd. Atendimentos","TMA","TME","Participação (%)"]] + rk[["Agente","Qtd_Atendimentos","TMA","TME","Participação_%"]].values.tolist()
        table_rank = Table(data_rank, hAlign="LEFT", colWidths=[65*mm,30*mm,35*mm,35*mm,30*mm])
        table_rank.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),0.4,colors.black),('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.whitesmoke,colors.lightgrey])]))
        elements.append(table_rank)
    else:
        elements.append(Paragraph("Sem dados.", styles['Small']))
    elements.append(PageBreak())

    # Heatmap
    elements.append(Paragraph("Heatmap – Volume (Dia × Hora)", styles['Heading2']))
    hm_buf = fig_heatmap_volume_image(df)
    if hm_buf is not None:
        elements.append(RLImage(hm_buf, width=170*mm, height=95*mm))
    else:
        elements.append(Paragraph("Sem dados suficientes para gerar o heatmap.", styles['Small']))

    elements.append(Spacer(1,8))
    elements.append(Paragraph(f"Relatório gerado em {BR_DATE(datetime.now())}", styles['Small']))
    doc.build(elements, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_word_relatorio(df, inicio_dt, fim_dt, thresholds_s, logo_bytes=None, signature_text=None):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    doc.add_heading("RELATÓRIO – ATENDIMENTO (CALL CENTER)", level=1)
    doc.add_paragraph(f"Empresa: Nevoli Telecom | Setor: Suporte Técnico / Atendimento")
    doc.add_paragraph(f"Período analisado: {BR_DATE(inicio_dt)} a {BR_DATE(fim_dt)}")
    doc.add_paragraph(" ")

    total, tma, tme = calc_kpis(df)
    tma_p = percentis_td(df["Tempo em Atendimento"]) if "Tempo em Atendimento" in df.columns else {}
    tme_p = percentis_td(df["Tempo em Espera"]) if "Tempo em Espera" in df.columns else {}

    doc.add_heading("KPIs do Período", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicador'
    hdr_cells[1].text = 'Valor'
    rows = [
        ("Total de Atendimentos", numero_br(total)),
        ("TMA Médio", format_timedelta(tma)),
        ("TME Médio (Espera)", format_timedelta(tme)),
        ("TMA p50/p75/p90/p95", f"{tma_p.get('p50','-')} / {tma_p.get('p75','-')} / {tma_p.get('p90','-')} / {tma_p.get('p95','-')}"),
        ("TME p50/p75/p90/p95", f"{tme_p.get('p50','-')} / {tme_p.get('p75','-')} / {tme_p.get('p90','-')} / {tme_p.get('p95','-')}")
    ]
    for r in rows:
        r_cells = table.add_row().cells
        r_cells[0].text = r[0]
        r_cells[1].text = str(r[1])

    doc.add_paragraph(" ")

    # Curva
    curva = curva_servico(df, thresholds_s)
    img_buf = fig_curva_servico_image(curva)
    doc.add_heading("Curva de Serviço (SLA de Espera)", level=2)
    try:
        doc.add_picture(img_buf, width=Inches(6.0))
    except Exception:
        try:
            tmp = io.BytesIO(img_buf.getvalue())
            doc.add_picture(tmp, width=Inches(6.0))
        except Exception:
            doc.add_paragraph("Imagem da curva indisponível.")
    doc.add_paragraph(" ")

    # Heatmap
    hm_buf = fig_heatmap_volume_image(df)
    doc.add_heading("Heatmap – Volume (Dia × Hora)", level=2)
    if hm_buf is not None:
        try:
            doc.add_picture(hm_buf, width=Inches(6.0))
        except Exception:
            try:
                tmp = io.BytesIO(hm_buf.getvalue())
                doc.add_picture(tmp, width=Inches(6.0))
            except Exception:
                doc.add_paragraph("Imagem do heatmap indisponível.")
    else:
        doc.add_paragraph("Sem dados suficientes para gerar o heatmap.")
    doc.add_paragraph(" ")

    # KPIs por agente segmentados (Atendimentos/TMA/TME/Participacao)
    doc.add_heading("KPIs por Agente (segmentado)", level=2)
    rk = tabela_rank(df).sort_values("Qtd_Atendimentos", ascending=False).reset_index(drop=True)

    # Atendimentos
    doc.add_heading("Atendimentos (maior → menor)", level=3)
    if not rk.empty:
        tbl = doc.add_table(rows=1, cols=2)
        hdr = tbl.rows[0].cells
        hdr[0].text = "Agente"; hdr[1].text = "Qtd Atendimentos"
        for _, r in rk[["Agente","Qtd_Atendimentos"]].iterrows():
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(r["Agente"]); row_cells[1].text = str(r["Qtd_Atendimentos"])
    else:
        doc.add_paragraph("Sem dados.")
    doc.add_paragraph(" ")

    # TMA
    doc.add_heading("TMA (médio) — menor → maior", level=3)
    if not rk.empty:
        rk_tma = rk.sort_values("TMA_td", ascending=True)
        tbl = doc.add_table(rows=1, cols=2); hdr = tbl.rows[0].cells; hdr[0].text="Agente"; hdr[1].text="TMA"
        for _, r in rk_tma[["Agente","TMA"]].iterrows():
            row_cells = tbl.add_row().cells; row_cells[0].text=str(r["Agente"]); row_cells[1].text=str(r["TMA"])
    else:
        doc.add_paragraph("Sem dados.")
    doc.add_paragraph(" ")

    # TME
    doc.add_heading("TME (médio) — menor → maior", level=3)
    if not rk.empty:
        rk_tme = rk.sort_values("TME_td", ascending=True)
        tbl = doc.add_table(rows=1, cols=2); hdr = tbl.rows[0].cells; hdr[0].text="Agente"; hdr[1].text="TME"
        for _, r in rk_tme[["Agente","TME"]].iterrows():
            row_cells = tbl.add_row().cells; row_cells[0].text=str(r["Agente"]); row_cells[1].text=str(r["TME"])
    else:
        doc.add_paragraph("Sem dados.")
    doc.add_paragraph(" ")

    # Participacao
    doc.add_heading("Participação (%)", level=3)
    if not rk.empty:
        tbl = doc.add_table(rows=1, cols=2); hdr = tbl.rows[0].cells; hdr[0].text="Agente"; hdr[1].text="Participação (%)"
        for _, r in rk[["Agente","Participação_%"]].iterrows():
            row_cells = tbl.add_row().cells; row_cells[0].text=str(r["Agente"]); row_cells[1].text=f"{r['Participação_%']:.1f}%"
    else:
        doc.add_paragraph("Sem dados.")
    doc.add_paragraph(" ")

    doc.add_paragraph(f"Relatório gerado em {BR_DATE(datetime.now())}")
    if signature_text:
        doc.add_paragraph(f"Assinado por: {signature_text}")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()

def export_excel_bytes(df_now, df_prev, rk_full, start_dt=None, end_dt=None):
    buf = io.BytesIO()
    sheet_now = f"Dados_{start_dt}" if start_dt else "Dados_Periodo"
    sheet_prev = f"Anterior_{start_dt}" if start_dt else "Periodo_Anterior"
    if rk_full is None:
        rk_full = pd.DataFrame()
    cols_essenciais = ["Agente","Equipe","Data Inicial","Hora Inicial","Tempo em Atendimento","Tempo em Espera"]
    df_now_small = df_now[[c for c in cols_essenciais if c in df_now.columns]].copy() if not df_now.empty else df_now.copy()
    df_prev_small = df_prev[[c for c in cols_essenciais if c in df_prev.columns]].copy() if not df_prev.empty else df_prev.copy()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df_now_small.to_excel(writer, sheet_name=sheet_now[:31], index=False)
        df_prev_small.to_excel(writer, sheet_name=sheet_prev[:31], index=False)
        if not rk_full.empty:
            rk_full.to_excel(writer, sheet_name="Ranking", index=False)
        summary = pd.DataFrame({"Métrica": ["Atendimentos_Periodo", "Atendimentos_Periodo_Anterior"], "Valor":[len(df_now), len(df_prev)]})
        summary.to_excel(writer, sheet_name="Resumo", index=False)
    buf.seek(0)
    return buf.getvalue()

def gerar_pdf_por_agente(agente, df_agent, inicio_dt, fim_dt, logo_bytes=None, signature_text=None):
    NumberedCanvas.logo_bytes = logo_bytes
    NumberedCanvas.signature_text = signature_text
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=18*mm, leftMargin=14*mm, rightMargin=14*mm)
    styles = getSampleStyleSheet(); styles.add(ParagraphStyle(name='Small', fontSize=9, leading=11))
    elements = []
    elements.append(Paragraph(f"RELATÓRIO INDIVIDUAL – {agente}", styles['Title'])); elements.append(Spacer(1,6))
    elements.append(Paragraph(f"Período: {BR_DATE(inicio_dt)} a {BR_DATE(fim_dt)}", styles['Normal'])); elements.append(Spacer(1,6))
    total, tma, tme = calc_kpis(df_agent)
    kpi_data = [["Métrica","Valor"],["Atendimentos", numero_br(total)],["TMA (médio)", format_timedelta(tma)],["TME (médio)", format_timedelta(tme)]]
    kpi_table = Table(kpi_data, hAlign="LEFT", colWidths=[60*mm,60*mm])
    kpi_table.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),0.5,colors.black),('FONTNAME',(0,0),(-1,0),'Helvetica-Bold')]))
    elements.append(kpi_table); elements.append(Spacer(1,6))
    elements.append(Paragraph("Últimos atendimentos (até 10)", styles['Heading2']))
    cols_show = [c for c in ["Data Inicial","Hora Inicial","Tempo em Atendimento","Tempo em Espera"] if c in df_agent.columns]
    if not df_agent.empty and cols_show:
        slice_table = df_agent.sort_values("Data Inicial", ascending=False).head(10)[cols_show]
        data_rows = [cols_show] + slice_table.values.tolist()
        t = Table(data_rows, hAlign="LEFT")
        t.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.4,colors.black),('BACKGROUND',(0,0),(-1,0),colors.lightgrey)]))
        elements.append(t)
    else:
        elements.append(Paragraph("Sem dados.", styles['Small']))
    elements.append(Spacer(1,8)); elements.append(Paragraph(f"Relatório gerado em {BR_DATE(datetime.now())}", styles['Small']))
    doc.build(elements, canvasmaker=NumberedCanvas); buffer.seek(0); return buffer.getvalue()

def gerar_zip_pdfs_por_agente(df_now, inicio_dt, fim_dt, logo_bytes=None, signature_text=None):
    agents = sorted(df_now['Agente'].dropna().unique()) if 'Agente' in df_now.columns else []
    in_mem_zip = io.BytesIO()
    with zipfile.ZipFile(in_mem_zip, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for ag in agents:
            df_agent = df_now[df_now['Agente'] == ag]
            pdf_bytes = gerar_pdf_por_agente(ag, df_agent, inicio_dt, fim_dt, logo_bytes=logo_bytes, signature_text=signature_text)
            filename = f"{ag.replace(' ','_')}_{inicio_dt}_{fim_dt}.pdf"
            zf.writestr(filename, pdf_bytes)
    in_mem_zip.seek(0); return in_mem_zip.getvalue()

# ------------------------------
# Paginação: callbacks e função de exibição (corrigida)
# ------------------------------
def _prev_page(key_prefix):
    st.session_state[key_prefix + '_page'] = max(0, st.session_state.get(key_prefix + '_page', 0) - 1)

def _next_page(key_prefix):
    pages = st.session_state.get(key_prefix + '_pages', 1)
    st.session_state[key_prefix + '_page'] = min(pages - 1, st.session_state.get(key_prefix + '_page', 0) + 1)

def show_paged_df(df, page_size=50, key_prefix='paged'):
    total = len(df)
    if total == 0:
        st.write("Sem linhas para exibir.")
        return
    pages = max(1, (total + page_size - 1) // page_size)
    # inicializa valores de sessão
    st.session_state.setdefault(key_prefix + '_page', 0)
    st.session_state[key_prefix + '_pages'] = pages

    p = st.session_state[key_prefix + '_page']
    start = p * page_size
    end = start + page_size

    col_left, col_center, col_right = st.columns([1,6,1])
    with col_left:
        st.button('◀️ Anterior', key=f'{key_prefix}_prev_btn', on_click=_prev_page, args=(key_prefix,))
    with col_right:
        st.button('Próxima ▶️', key=f'{key_prefix}_next_btn', on_click=_next_page, args=(key_prefix,))

    st.write(f"Página {p+1} de {pages} — mostrando linhas {start+1} a {min(end,total)} de {total}")
    st.dataframe(df.iloc[start:end], use_container_width=True)

# ------------------------------
# UI - Sidebar
# ------------------------------
st.sidebar.title("📊 Navegação")
pagina = st.sidebar.radio("Ir para:", ["Geral", "Individual", "Ranking", "Relatório"])
uploaded_file = st.sidebar.file_uploader("📂 Envie o CSV do Chat", type=["csv"])

with st.sidebar.expander("Configurações avançadas", expanded=False):
    logo_file = st.file_uploader("Logo (PNG/JPG) - opcional", type=["png","jpg","jpeg"])
    signature_input = st.text_input("Assinatura (aparece no rodapé do PDF/Word)", value="")
    excluir_bot = st.checkbox("Excluir bot (SYNTESIS - Olivia Bot)", value=True)
    st.markdown("**Limiares TMA (em minutos)**")
    tma_warn_min = st.number_input("TMA aviso — amarelo (min)", min_value=0, value=50)
    tma_crit_min = st.number_input("TMA crítico — vermelho (min)", min_value=0, value=60)
    tma_warn_s = int(tma_warn_min * 60)
    tma_crit_s = int(tma_crit_min * 60)
    st.markdown("---")
    sla_target = st.number_input("SLA alvo (segundos)", min_value=1, value=5)
    sla_green_pct = st.slider("SLA verde — % mínimo", min_value=0, max_value=100, value=80)
    sla_yellow_pct = st.slider("SLA amarelo — % mínimo", min_value=0, max_value=100, value=60)
    sla_thresholds = st.multiselect("Alvos de SLA (s) - curva", options=[5,10,20,30,60], default=[5])
    st.markdown("Deduplicação")
    st.markdown("---")
    page_size = st.number_input("Linhas por página (visualização)", min_value=10, max_value=1000, value=50)

if not uploaded_file:
    st.title("Dashboard Call Center – Nevoli")
    st.write("Envie o CSV do chat na barra lateral. Use 'Configurações avançadas' para ajustar limites e deduplicação.")
    st.stop()

# ------------------------------
# Processar arquivo (ler uma vez)
# ------------------------------
try:
    with st.spinner('Lendo CSV...'):
        uploaded_file.seek(0)
        df_raw = read_csv_bytes(uploaded_file)
        st.session_state['df_raw_orig'] = df_raw.copy()
except Exception as e:
    st.error(f"Erro ao ler CSV: {e}")
    st.stop()

# dedupe choice
id_col_choice = None
try:
    id_options = df_raw.columns.tolist()
    id_options.insert(0, "Nenhuma / fallback")
    sel = st.sidebar.selectbox("Coluna de ID para dedupe (opcional)", options=id_options, index=0)
    if sel and sel != "Nenhuma / fallback":
        id_col_choice = sel
except Exception:
    id_col_choice = None

with st.spinner('Preparando dados...'):
    df = preparacao_base(df_raw, excluir_bot=True if 'excluir_bot' not in locals() else excluir_bot, id_col=id_col_choice)
removed = df.attrs.get('duplicates_removed', 0)
dedupe_id_col = df.attrs.get('dedupe_id_col', None)
if removed:
    st.sidebar.warning(f"{removed} linhas removidas por deduplicação ({'col: '+str(dedupe_id_col) if dedupe_id_col else 'fallback'})")

# periodo presets
min_dt = df['Data Inicial'].min().date() if 'Data Inicial' in df.columns and not df['Data Inicial'].isna().all() else date.today()-timedelta(days=7)
max_dt = df['Data Inicial'].max().date() if 'Data Inicial' in df.columns and not df['Data Inicial'].isna().all() else date.today()

preset = st.sidebar.selectbox('Período rápido', options=['Personalizado', 'Último mês', 'Mês atual', 'Últimos 30 dias'], index=0)
if preset == 'Último mês':
    today = date.today()
    first = (today.replace(day=1) - timedelta(days=1)).replace(day=1)
    last = today.replace(day=1) - timedelta(days=1)
    preset_start, preset_end = first, last
elif preset == 'Mês atual':
    today = date.today()
    preset_start = today.replace(day=1)
    preset_end = max_dt
elif preset == 'Últimos 30 dias':
    preset_end = max_dt
    preset_start = max_dt - timedelta(days=29)
else:
    preset_start, preset_end = min_dt, max_dt

start_dt, end_dt = st.sidebar.date_input("Período", (preset_start, preset_end))
if isinstance(start_dt, tuple):
    start_dt, end_dt = start_dt[0], start_dt[1]

# ------------------------------
# aplicar filtros: apenas periodo
# ------------------------------
df_now, df_prev, prev_start, prev_end = split_periodos(df, start_dt, end_dt)
df_filtrado = df_now.copy()

# remover categorias não usadas
if 'Agente' in df_filtrado.columns and pd.api.types.is_categorical_dtype(df_filtrado['Agente']):
    try:
        df_filtrado['Agente'] = df_filtrado['Agente'].cat.remove_unused_categories()
    except Exception:
        pass

# preparar logo
logo_bytes = None
if 'logo_file' in locals() and logo_file:
    try:
        uploaded_bytes = logo_file.read()
        buf = io.BytesIO(uploaded_bytes)
        img = PILImage.open(buf); img = img.convert("RGBA")
        out = io.BytesIO(); img.save(out, format="PNG"); logo_bytes = out.getvalue(); out.close(); buf.close()
    except Exception:
        logo_bytes = None

# ------------------------------
# PÁGINAS
# ------------------------------
if pagina == "Geral":
    st.title("📈 Dashboard Geral")
    total_now, tma_now, tme_now = calc_kpis(df_filtrado)
    total_prev, tma_prev, tme_prev = calc_kpis(df_prev)
    var_total = ((total_now - total_prev)/total_prev*100) if total_prev else None
    col1, col2, col3 = st.columns(3)
    col1.metric("📞 Atendimentos (filtrado)", numero_br(total_now), f"{var_total:+.1f}%" if var_total is not None else "—")
    col2.metric("⏱️ TMA (médio)", format_timedelta(tma_now))
    col3.metric("⌛ TME (médio)", format_timedelta(tme_now))
    st.caption(f"Período: {BR_DATE(start_dt)} → {BR_DATE(end_dt)}")

    if "Tempo em Espera" in df_filtrado.columns and not df_filtrado.empty:
        overall_sla_pct = (df_filtrado["Tempo em Espera"].dt.total_seconds() <= sla_target).mean() * 100
        sla_color = "🟢" if overall_sla_pct >= sla_green_pct else ("🟡" if overall_sla_pct >= sla_yellow_pct else "🔴")
        st.metric(f"SLA ≤ {sla_target}s", f"{overall_sla_pct:.1f}%", sla_color)

    colA, colB = st.columns([1,1])
    with colA:
        st.plotly_chart(plot_curva_servico(df_filtrado, sla_thresholds), use_container_width=True)
    with colB:
        tma_p = percentis_td(df_filtrado["TMA_s"]) if "TMA_s" in df_filtrado.columns else percentis_td(df_filtrado["Tempo em Atendimento"]) if "Tempo em Atendimento" in df_filtrado.columns else {}
        tme_p = percentis_td(df_filtrado["TME_s"]) if "TME_s" in df_filtrado.columns else percentis_td(df_filtrado["Tempo em Espera"]) if "Tempo em Espera" in df_filtrado.columns else {}
        st.write("**TMA p50/p75/p90/p95**"); st.write(f"{tma_p.get('p50','-')} / {tma_p.get('p75','-')} / {tma_p.get('p90','-')} / {tma_p.get('p95','-')}")
        st.write("**TME p50/p75/p90/p95**"); st.write(f"{tme_p.get('p50','-')} / {tme_p.get('p75','-')} / {tme_p.get('p90','-')} / {tme_p.get('p95','-')}")

    if 'Agente' in df_filtrado.columns and not df_filtrado.empty:
        dist = df_filtrado.groupby("Agente").size().reset_index(name="Qtd_Atendimentos").sort_values("Qtd_Atendimentos", ascending=False)
        fig = px.bar(dist.head(50), x="Agente", y="Qtd_Atendimentos", text="Qtd_Atendimentos", height=380)
        fig.update_traces(textposition='outside'); fig.update_layout(xaxis_tickangle=-45, template="plotly_white")
        st.plotly_chart(fig, use_container_width=True)

    heat_fig = plot_heatmap(df_filtrado)
    if heat_fig:
        st.plotly_chart(heat_fig, use_container_width=True)

    st.write("### 📋 Detalhamento (filtrado)")
    cols_show = [c for c in ["Agente","Data Inicial","Hora Inicial","Tempo em Atendimento","Tempo em Espera","Mes_Ano"] if c in df_filtrado.columns]
    if cols_show:
        show_paged_df(df_filtrado[cols_show].sort_values("Data Inicial", ascending=False), page_size=page_size, key_prefix='geral')

elif pagina == "Individual":
    st.title("👤 Comparativo Individual (até 3 agentes)")
    agentes = sorted(list(df_filtrado['Agente'].dropna().unique())) if 'Agente' in df_filtrado.columns else []
    selecionados = st.multiselect("Selecione até 3 agentes", options=agentes, default=agentes[:1])
    if len(selecionados) > 3:
        st.warning("Máx 3 agentes — considerando os 3 primeiros.")
        selecionados = selecionados[:3]

    if selecionados:
        base_sel = df_filtrado[df_filtrado["Agente"].isin(selecionados)]
        for ag in selecionados:
            sub = base_sel[base_sel["Agente"] == ag]
            total, tma, tme = calc_kpis(sub)
            st.subheader(f"📊 {ag}")
            c1, c2, c3 = st.columns(3)
            c1.metric("Atendimentos", numero_br(total))
            c2.metric("TMA", format_timedelta(tma))
            c3.metric("TME", format_timedelta(tme))
            cols_show = [c for c in ["Data Inicial","Hora Inicial","Tempo em Atendimento","Tempo em Espera","Mes_Ano"] if c in sub.columns]
            if cols_show:
                show_paged_df(sub[cols_show].sort_values("Data Inicial", ascending=False), page_size=page_size, key_prefix=f'ind_{ag}')

elif pagina == "Ranking":
    st.title("🏆 Ranking dos Agentes")
    rk = tabela_rank(df_filtrado)
    if not rk.empty:
        rk = rk.sort_values("Qtd_Atendimentos", ascending=False).reset_index(drop=True)
    outs_tma = detectar_outliers_iqr(rk, "TMA_td") if not rk.empty else set()
    outs_tme = detectar_outliers_iqr(rk, "TME_td") if not rk.empty else set()

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("👥 Agentes", rk.shape[0] if not rk.empty else 0)
    c2.metric("📞 Total Atendimentos (filtrado)", numero_br(rk["Qtd_Atendimentos"].sum()) if not rk.empty else "0")
    try:
        tma_med = pd.to_timedelta(rk['TMA_td'].dropna().apply(lambda x: x.total_seconds()).mean(), unit='s') if 'TMA_td' in rk.columns and not rk['TMA_td'].isna().all() else (df_filtrado["Tempo em Atendimento"].mean() if "Tempo em Atendimento" in df_filtrado.columns else pd.Timedelta(0))
    except Exception:
        tma_med = pd.Timedelta(0)
    try:
        tme_med = pd.to_timedelta(rk['TME_td'].dropna().apply(lambda x: x.total_seconds()).mean(), unit='s') if 'TME_td' in rk.columns and not rk['TME_td'].isna().all() else (df_filtrado["Tempo em Espera"].mean() if "Tempo em Espera" in df_filtrado.columns else pd.Timedelta(0))
    except Exception:
        tme_med = pd.Timedelta(0)
    c3.metric("⏱️ TMA Médio", format_timedelta(tma_med))
    c4.metric("⌛ TME Médio", format_timedelta(tme_med))
    st.caption(f"Período: {BR_DATE(start_dt)} → {BR_DATE(end_dt)}")

    if not rk.empty:
        try:
            tma_warn_s
            tma_crit_s
        except NameError:
            tma_warn_s = 5 * 60; tma_crit_s = 10 * 60

        status_list = []
        for _, row in rk.iterrows():
            ag = row["Agente"]
            df_ag = df_filtrado[df_filtrado["Agente"] == ag]
            sla_pct = (df_ag["TME_s"] <= sla_target).mean() * 100 if ("TME_s" in df_ag.columns and not df_ag.empty) else None
            tma_sec = row["TMA_td"].total_seconds() if pd.notna(row["TMA_td"]) else None
            status = "⚪"
            if sla_pct is not None:
                if sla_pct >= sla_green_pct:
                    status = "🟢"
                elif sla_pct >= sla_yellow_pct:
                    status = "🟡"
                else:
                    status = "🔴"
            if tma_sec is not None:
                if tma_sec > tma_crit_s:
                    status = "🔴"
                elif tma_sec > tma_warn_s and status != "🔴":
                    status = "🟡"
            status_list.append(status)

        rk_display = rk.copy()
        rk_display["Status"] = status_list
        rk_display["Badge"] = ""
        rk_display.loc[rk_display.head(3).index, "Badge"] = "🏅 Top 3"
        rk_display.loc[rk_display["Agente"].isin(outs_tma | outs_tme), "Badge"] = rk_display["Badge"].apply(lambda x: (x + " ⚠️").strip())

        show_paged_df(rk_display[["Agente","Qtd_Atendimentos","Participação_%","TMA","TME","Status","Badge"]], page_size=page_size, key_prefix='ranking')
    else:
        st.info("Sem dados no período / filtro selecionado.")

elif pagina == "Relatório":
    st.title("📄 Relatório")
    st.caption(f"Período: {BR_DATE(start_dt)} → {BR_DATE(end_dt)}")
    st.write("Relatório com KPIs, Curva SLA, KPIs por Agente segmentados, Ranking e Heatmap. Gera PDF e Word editável.")

    if st.button("📥 Gerar PDF (com assinatura)"):
        thresholds = sla_thresholds if 'sla_thresholds' in locals() else [5,10,20]
        with st.spinner('Gerando PDF...'):
            pdf_bytes = gerar_pdf_relatorio(df_filtrado, datetime.combine(start_dt, datetime.min.time()), datetime.combine(end_dt, datetime.max.time()), thresholds, logo_bytes=logo_bytes, signature_text=signature_input or None)
        st.download_button("Download PDF", pdf_bytes, file_name=f"relatorio_{start_dt}_{end_dt}.pdf", mime="application/pdf")

    if st.button("📄 Gerar Word (.docx) editável"):
        thresholds = sla_thresholds if 'sla_thresholds' in locals() else [5,10,20]
        with st.spinner('Gerando Word...'):
            word_bytes = gerar_word_relatorio(df_filtrado, datetime.combine(start_dt, datetime.min.time()), datetime.combine(end_dt, datetime.max.time()), thresholds, logo_bytes=logo_bytes, signature_text=signature_input or None)
        st.download_button("Download Word (.docx)", word_bytes, file_name=f"relatorio_{start_dt}_{end_dt}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if st.button("📎 Gerar PDFs por Agente (ZIP)"):
        with st.spinner('Gerando PDFs por agente...'):
            zip_bytes = gerar_zip_pdfs_por_agente(df_filtrado, start_dt, end_dt, logo_bytes=logo_bytes, signature_text=signature_input or None)
        st.download_button("Download ZIP", zip_bytes, file_name=f"pdfs_agentes_{start_dt}_{end_dt}.zip", mime="application/zip")

    if st.button("⬇️ Exportar Excel (Dados + Ranking)"):
        rk_full = tabela_rank(df_filtrado)
        with st.spinner('Gerando Excel...'):
            excel_bytes = export_excel_bytes(df_filtrado, df_prev, rk_full, start_dt=start_dt, end_dt=end_dt)
        st.download_button("Download Excel", excel_bytes, file_name=f"relatorio_{start_dt}_{end_dt}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

st.markdown("---")
st.caption("Dica: nomes de agente são limpos automaticamente (removemos emojis, BOMs e nomes inválidos).")
